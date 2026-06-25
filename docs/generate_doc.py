from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# PORTADA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FINTRACK')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 212, 170)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentación Técnica de Refactorización')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(200, 200, 200)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nArquitectura Hexagonal + Principios SOLID\n')
run.font.size = Pt(14)
run.font.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\nDocumento generado: Junio 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDO
# ============================================================
doc.add_heading('Índice', level=1)
toc_items = [
    '1. Resumen Ejecutivo',
    '2. Arquitectura Original vs. Hexagonal',
    '3. Principios SOLID Implementados',
    '4. Estructura de Directorios',
    '5. Backend: Domain Layer',
    '6. Backend: Application Layer (Casos de Uso)',
    '7. Backend: Infrastructure Layer',
    '8. Backend: DI Container',
    '9. Frontend: Separación de Responsabilidades',
    '10. Mapeo de Endpoints',
    '11. Cómo Ejecutar',
    '12. Conclusión',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ============================================================
# 1. RESUMEN EJECUTIVO
# ============================================================
doc.add_heading('1. Resumen Ejecutivo', level=1)
doc.add_paragraph(
    'FinTrack es un sistema de gestión financiera personal compuesto por un backend '
    'en TypeScript (Hono.js + SQLite) y un frontend en Python (Streamlit). '
    'El proyecto fue refactorizado completamente para adoptar una Arquitectura Hexagonal '
    '(también conocida como Ports & Adapters) y cumplir con los cinco principios SOLID.'
)
doc.add_paragraph(
    'La refactorización se realizó sin alterar la funcionalidad existente: todos los '
    'endpoints de la API, las consultas y las interfaces de usuario se mantienen idénticas. '
    'El cambio es puramente estructural y de organización del código.'
)

# ============================================================
# 2. ARQUITECTURA ORIGINAL VS HEXAGONAL
# ============================================================
doc.add_heading('2. Arquitectura Original vs. Hexagonal', level=1)

doc.add_heading('2.1 Arquitectura Original (Layered)', level=2)
doc.add_paragraph(
    'El backend original seguía una arquitectura de 2 capas simple:'
)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Capa', 'Archivos', 'Problemas']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['Routes', 'src/routes/*.ts', 'Mezclaban HTTP + lógica de negocio'],
    ['Data Access', 'src/lib/db.ts (375 líneas)', 'GOD OBJECT: CRUD + stats + exchange rates en 1 archivo'],
    ['DB Config', 'src/lib/database.ts', 'Conexión + esquema (correcto)'],
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph()
doc.add_heading('2.2 Arquitectura Hexagonal (Nueva)', level=2)
doc.add_paragraph(
    'La nueva arquitectura separa el código en tres anillos concéntricos:'
)

items = [
    ('Domain (Núcleo)', 'Entidades, Value Objects, interfaces de repositorio (puertos), servicios de dominio. Sin dependencias externas.'),
    ('Application', 'Casos de uso que orquestan la lógica de dominio. Dependen solo de interfaces (puertos).'),
    ('Infrastructure', 'Implementaciones concretas: repositorios SQLite, controladores HTTP, cliente API externa.'),
]
for title, desc in items:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    'El flujo de datos sigue una dirección única: '
    'HTTP Controller → Use Case → Repository Interface ← SQLite Repository. '
    'Nunca al revés (Dependency Inversion).'
)

doc.add_page_break()

# ============================================================
# 3. PRINCIPIOS SOLID
# ============================================================
doc.add_heading('3. Principios SOLID Implementados', level=1)

solids = [
    ('S - Single Responsibility Principle',
     'Cada archivo tiene una única responsabilidad:\n'
     '• Cada entidad en su propio archivo (Transaction.ts, Budget.ts, Goal.ts)\n'
     '• Cada caso de uso en su propio archivo (CreateTransactionUseCase.ts, etc.)\n'
     '• Cada repositorio en su propio archivo (SqliteTransactionRepository.ts, etc.)\n'
     '• El antiguo db.ts (375 líneas) fue eliminado como monolito; su lógica se distribuyó en ~20 archivos especializados.'),
    ('O - Open/Closed Principle',
     'Las entidades y los casos de uso están abiertos a extensión pero cerrados a modificación.\n'
     'Para añadir una nueva entidad (ej. "Inversiones"):\n'
     '1. Crear Inversión.ts en domain/entities/\n'
     '2. Crear IInvestmentRepository.ts en domain/repositories/\n'
     '3. Crear casos de uso en application/investments/\n'
     '4. Crear SqliteInvestmentRepository.ts en infrastructure/persistence/\n'
     'Sin modificar ningún archivo existente.'),
    ('L - Liskov Substitution Principle',
     'Las interfaces de repositorio (ITransactionRepository, etc.) garantizan que cualquier '
     'implementación (SQLite, PostgreSQL mock, etc.) puede sustituirse sin afectar a los consumidores. '
     'Actualmente solo existe Sqlite* pero el diseño permite agregar otras sin cambios.'),
    ('I - Interface Segregation Principle',
     'Cada caso de uso recibe solo la interfaz que necesita. Por ejemplo, CreateTransactionUseCase '
     'recibe ITransactionRepository y solo usa create(), no tiene acceso a delete() o getAll(). '
     'Antes, todos los consumers importaban todo db.ts con 20+ funciones.'),
    ('D - Dependency Inversion Principle',
     'Las capas superiores (controllers) dependen de abstracciones (interfaces), no de implementaciones concretas.\n'
     'El DI Container en di/container.ts construye toda la cadena de dependencias:\n'
     '• Repositorios concretos → inyectados en Casos de Uso\n'
     '• Casos de Uso → inyectados en Controladores\n'
     '• Controladores → montados en rutas de Hono'),
]
for title, desc in solids:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# 4. ESTRUCTURA DE DIRECTORIOS
# ============================================================
doc.add_heading('4. Estructura de Directorios', level=1)

code_block = """
fintrack-api/src/
├── index.ts                     # Punto de entrada (solo DI + server)
├── openapi.ts                   # Spec OpenAPI (sin cambios)
│
├── domain/                      # NÚCLEO (sin dependencias)
│   ├── entities/
│   │   ├── Transaction.ts
│   │   ├── Budget.ts
│   │   ├── Goal.ts
│   │   └── ExchangeRate.ts
│   ├── repositories/            # Puertos (interfaces)
│   │   ├── ITransactionRepository.ts
│   │   ├── IBudgetRepository.ts
│   │   ├── IGoalRepository.ts
│   │   └── IExchangeRateRepository.ts
│   └── services/                # Lógica de negocio pura
│       ├── StatsService.ts
│       └── BudgetStatusService.ts
│
├── application/                 # Casos de uso
│   ├── transactions/
│   │   ├── GetTransactionsUseCase.ts
│   │   ├── CreateTransactionUseCase.ts
│   │   ├── UpdateTransactionUseCase.ts
│   │   ├── DeleteTransactionUseCase.ts
│   │   └── BulkDeleteTransactionsUseCase.ts
│   ├── budgets/
│   │   ├── GetBudgetsUseCase.ts
│   │   ├── CreateBudgetUseCase.ts
│   │   ├── UpdateBudgetUseCase.ts
│   │   ├── DeleteBudgetUseCase.ts
│   │   ├── BulkDeleteBudgetsUseCase.ts
│   │   └── GetBudgetStatusUseCase.ts
│   ├── goals/
│   │   ├── GetGoalsUseCase.ts
│   │   ├── CreateGoalUseCase.ts
│   │   ├── DepositToGoalUseCase.ts
│   │   ├── DeleteGoalUseCase.ts
│   │   └── BulkDeleteGoalsUseCase.ts
│   └── exchange-rates/
│       ├── GetExchangeRatesUseCase.ts
│       └── RefreshExchangeRatesUseCase.ts
│
├── infrastructure/              # Adaptadores
│   ├── ExchangeRateApi.ts       # Cliente HTTP externo
│   ├── persistence/
│   │   ├── SqliteTransactionRepository.ts
│   │   ├── SqliteBudgetRepository.ts
│   │   ├── SqliteGoalRepository.ts
│   │   └── SqliteExchangeRateRepository.ts
│   └── http/controllers/
│       ├── TransactionController.ts
│       ├── BudgetController.ts
│       ├── GoalController.ts
│       ├── StatsController.ts
│       └── ExchangeRateController.ts
│
├── di/
│   └── container.ts             # Inyección de dependencias
│
├── lib/                         # Archivos legacy mantenidos
│   ├── db.ts                    # Ya no se importa desde index.ts
│   ├── database.ts              # Conexión SQLite (reutilizada)
│   └── validation.ts            # Schemas Zod (sin cambios)
│
└── scripts/
    └── migrate.ts
"""
p = doc.add_paragraph()
run = p.add_run(code_block)
run.font.size = Pt(8)
run.font.name = 'Consolas'

doc.add_page_break()

# ============================================================
# 5. BACKEND DOMAIN
# ============================================================
doc.add_heading('5. Backend: Domain Layer', level=1)
doc.add_paragraph(
    'El dominio es el corazón de la aplicación. No importa nada de infraestructura. '
    'Contiene:'
)
doc.add_heading('5.1 Entidades', level=2)
doc.add_paragraph(
    'Transaction, Budget, Goal y ExchangeRate son interfaces TypeScript puras '
    'que definen la forma de los datos sin comportamiento. No tienen decoradores, '
    'no extienden clases de framework, no dependen de la base de datos.'
)

doc.add_heading('5.2 Puertos (Interfaces de Repositorio)', level=2)
doc.add_paragraph(
    'I*Repository define los contratos que la infraestructura debe implementar:'
)
items = [
    'ITransactionRepository: getAll, getById, create, update, delete',
    'IBudgetRepository: getAll, getById, create, update, delete',
    'IGoalRepository: getAll, getById, create, deposit, delete',
    'IExchangeRateRepository: getRates, getRate, upsertRates, isCacheFresh',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Servicios de Dominio', level=2)
doc.add_paragraph(
    'StatsService contiene toda la lógica de estadísticas: resumen mensual, '
    'gastos por categoría, tendencias, top gastos y mapa de calor. '
    'Recibe ITransactionRepository por constructor y opera sobre los datos en memoria.'
)
doc.add_paragraph(
    'BudgetStatusService cruza presupuestos con transacciones para calcular '
    'porcentajes de consumo. Depende de IBudgetRepository e ITransactionRepository.'
)

# ============================================================
# 6. APPLICATION LAYER
# ============================================================
doc.add_heading('6. Backend: Application Layer (Casos de Uso)', level=1)
doc.add_paragraph(
    'Cada caso de uso es una clase con un único método execute(). Representa '
    'una operación atómica del sistema. La inyección de dependencias se hace '
    'por constructor.'
)
doc.add_paragraph('Ejemplo típico:')
p = doc.add_paragraph()
run = p.add_run(
    'class CreateTransactionUseCase {\n'
    '  constructor(private readonly txRepo: ITransactionRepository) {}\n\n'
    '  execute(data: Omit<Transaction, "id">): Transaction {\n'
    '    return this.txRepo.create(data);\n'
    '  }\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph('Casos de uso implementados:')
use_cases = [
    'Transactions: Get, Create, Update, Delete, BulkDelete',
    'Budgets: Get, Create, Update, Delete, BulkDelete, GetStatus',
    'Goals: Get, Create, Deposit, Delete, BulkDelete',
    'Exchange Rates: Get, Refresh',
]
for uc in use_cases:
    doc.add_paragraph(uc, style='List Bullet')

doc.add_page_break()

# ============================================================
# 7. INFRASTRUCTURE
# ============================================================
doc.add_heading('7. Backend: Infrastructure Layer', level=1)

doc.add_heading('7.1 Persistencia (SQLite)', level=2)
doc.add_paragraph(
    'Cuatro clases que implementan las interfaces del dominio:'
)
items = [
    'SqliteTransactionRepository → ITransactionRepository',
    'SqliteBudgetRepository → IBudgetRepository',
    'SqliteGoalRepository → IGoalRepository',
    'SqliteExchangeRateRepository → IExchangeRateRepository',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')
doc.add_paragraph(
    'Todas usan better-sqlite3 y reutilizan getDb() desde src/lib/database.ts '
    '(sin cambios respecto al original).'
)

doc.add_heading('7.2 Controladores HTTP', level=2)
doc.add_paragraph(
    'Los controladores reciben casos de uso por constructor. Son ultradelgados:\n'
    '• Parsean el request HTTP\n'
    '• Validan con Zod (schemas desde lib/validation.ts)\n'
    '• Ejecutan el caso de uso\n'
    '• Devuelven la respuesta JSON\n\n'
    'NO contienen lógica de negocio ni acceso a base de datos.'
)

doc.add_heading('7.3 ExchangeRateApi', level=2)
doc.add_paragraph(
    'Adaptador para la API externa open.er-api.com. Encapsula la lógica de fetch '
    'y parseo de JSON. Si en el futuro se cambia de proveedor, solo se modifica esta clase.'
)

# ============================================================
# 8. DI CONTAINER
# ============================================================
doc.add_heading('8. DI Container', level=1)
doc.add_paragraph(
    'El archivo di/container.ts contiene la función buildContainer() que '
    'construye todo el grafo de dependencias:'
)
p = doc.add_paragraph()
run = p.add_run(
    'export function buildContainer(): Container {\n'
    '  const txRepo = new SqliteTransactionRepository();\n'
    '  const budgetRepo = new SqliteBudgetRepository();\n'
    '  // ... más repositorios\n\n'
    '  const statsService = new StatsService(txRepo);\n\n'
    '  const createTx = new CreateTransactionUseCase(txRepo);\n'
    '  // ... más casos de uso\n\n'
    '  return {\n'
    '    transactionController: new TransactionController(getTxs, createTx, ...),\n'
    '    // ... más controladores\n'
    '  };\n'
    '}'
)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph(
    'El index.ts simplificado importa buildContainer(), monta las rutas '
    'asociando cada controller a su path, y arranca el servidor. El resultado '
    'es un 70% más corto que el original.'
)

doc.add_page_break()

# ============================================================
# 9. FRONTEND
# ============================================================
doc.add_heading('9. Frontend: Separación de Responsabilidades', level=1)
doc.add_paragraph(
    'El frontend en Streamlit fue refactorizado para separar la lógica en módulos '
    'especializados. El antiguo ui_tweak.py (258 líneas, GOD OBJECT) ahora es '
    'una capa de compatibilidad delgada (~15 líneas activas).'
)

doc.add_heading('9.1 Servicios', level=2)
services = [
    ('services/api_client.py', 'Cliente HTTP centralizado. Cada página usaba requests.get hardcodeado con API_URL repetido 15+ veces. Ahora hay un solo ApiClient con métodos named: get_transactions(), create_budget(), etc.'),
    ('services/currency_service.py', 'Maneja la moneda seleccionada, el archivo user_preferences.json, y la conversión con tasa de cambio. Incluye fmt_money() y fmt_html_money().'),
    ('services/cache_service.py', 'Caché SQLite local para tasas de cambio. Antes estaba mezclado en ui_tweak.py con init_cache(), get_cached_rate_from_local(), cache_rates_locally().'),
]
for title, desc in services:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('9.2 Modelos', level=2)
doc.add_paragraph(
    'models/models.py define dataclasses Python (Transaction, Budget, BudgetStatus, Goal, MonthlySummary) '
    'que tipan los datos que viajan entre frontend y backend. Antes no había tipado explícito.'
)

doc.add_heading('9.3 UI Components', level=2)
ui_comps = [
    ('ui/styles.py', 'Todo el CSS global movido aquí. Antes estaba esparcido entre Dashboard.py y ui_tweak.py.'),
    ('ui/components.py', 'Componentes reutilizables: render_kpi_card(), render_donut_chart(), render_trend_chart(), render_top_expenses(). Invocados desde Dashboard.py.'),
    ('ui/navigation.py', 'Sidebar: render_sidebar(), _render_connection_status(), _render_currency_selector().'),
]
for title, desc in ui_comps:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('9.4 Páginas', level=2)
doc.add_paragraph(
    'Las 5 páginas (Dashboard.py + pages/*) se simplificaron porque ahora '
    'obtienen el ApiClient desde ui_tweak.get_api_client() en lugar de '
    'hardcodear API_URL y hacer requests directamente. '
    'La lógica de UI se mantiene en las páginas, pero la lógica de datos '
    'y comunicación está en los servicios.'
)

doc.add_page_break()

# ============================================================
# 10. MAPEO DE ENDPOINTS
# ============================================================
doc.add_heading('10. Mapeo de Endpoints', level=1)
doc.add_paragraph('Todos los endpoints se mantienen exactamente igual que antes de la refactorización:')

endpoints = [
    ('GET', '/transactions', 'TransactionController.getAll'),
    ('POST', '/transactions', 'TransactionController.create'),
    ('PATCH', '/transactions/:id', 'TransactionController.update'),
    ('DELETE', '/transactions/:id', 'TransactionController.delete'),
    ('POST', '/transactions/bulk-delete', 'TransactionController.bulkDelete'),
    ('GET', '/stats/summary', 'StatsController.summary'),
    ('GET', '/stats/by-category', 'StatsController.byCategory'),
    ('GET', '/stats/trends', 'StatsController.trends'),
    ('GET', '/stats/top-expenses', 'StatsController.topExpenses'),
    ('GET', '/stats/heatmap', 'StatsController.heatmap'),
    ('GET', '/budgets', 'BudgetController.getAll'),
    ('GET', '/budgets/status', 'BudgetController.getStatus'),
    ('POST', '/budgets', 'BudgetController.create'),
    ('PATCH', '/budgets/:id', 'BudgetController.update'),
    ('DELETE', '/budgets/:id', 'BudgetController.delete'),
    ('POST', '/budgets/bulk-delete', 'BudgetController.bulkDelete'),
    ('GET', '/goals', 'GoalController.getAll'),
    ('POST', '/goals', 'GoalController.create'),
    ('PATCH', '/goals/:id/deposit', 'GoalController.deposit'),
    ('DELETE', '/goals/:id', 'GoalController.delete'),
    ('POST', '/goals/bulk-delete', 'GoalController.bulkDelete'),
    ('GET', '/exchange-rates', 'ExchangeRateController.getAll'),
    ('GET', '/exchange-rates/:currency', 'ExchangeRateController.getByCurrency'),
    ('POST', '/exchange-rates/refresh', 'ExchangeRateController.refresh'),
]
table = doc.add_table(rows=len(endpoints) + 1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Método', 'Endpoint', 'Controlador']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

for r, (method, path, controller) in enumerate(endpoints, 1):
    table.rows[r].cells[0].text = method
    table.rows[r].cells[1].text = path
    table.rows[r].cells[2].text = controller

doc.add_page_break()

# ============================================================
# 11. CÓMO EJECUTAR
# ============================================================
doc.add_heading('11. Cómo Ejecutar', level=1)

doc.add_heading('11.1 Backend (API)', level=2)
doc.add_paragraph('Requisitos: Node.js 20+, npm')
p = doc.add_paragraph()
run = p.add_run(
    'cd fintrack-api\n'
    'npm install          # solo la primera vez\n'
    'npm run migrate      # migrar datos JSON → SQLite (opcional)\n'
    'npm run dev          # desarrollo con hot reload\n'
    '# o: npm start       # producción'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('La API corre en http://localhost:3000')
doc.add_paragraph('Swagger UI: http://localhost:3000/swagger')

doc.add_heading('11.2 Frontend (Dashboard)', level=2)
doc.add_paragraph('Requisitos: Python 3.10+, pip')
p = doc.add_paragraph()
run = p.add_run(
    'cd fintrack-dashboard\n'
    'pip install -r requirements.txt   # solo la primera vez\n'
    'streamlit run Dashboard.py'
)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('El dashboard abre en http://localhost:8501')

doc.add_heading('11.3 Orden de Arranque', level=2)
doc.add_paragraph('1. Iniciar el backend (npm run dev en fintrack-api/)')
doc.add_paragraph('2. Iniciar el frontend (streamlit run Dashboard.py en fintrack-dashboard/)')
doc.add_paragraph('3. Navegar a http://localhost:8501')

doc.add_page_break()

# ============================================================
# 12. CONCLUSIÓN
# ============================================================
doc.add_heading('12. Conclusión', level=1)
doc.add_paragraph(
    'La refactorización transformó FinTrack de una arquitectura monolítica de 2 capas '
    'a una arquitectura hexagonal limpia que respeta los 5 principios SOLID. '
    'Los beneficios clave son:'
)

benefits = [
    'Mantenibilidad: Cada archivo tiene una responsabilidad única. Localizar y corregir bugs es más rápido.',
    'Testeabilidad: Los casos de uso y servicios de dominio pueden probarse con mocks de repositorios.',
    'Extensibilidad: Añadir nuevas entidades no requiere modificar código existente (OCP).',
    'Separación de Concerns: La lógica de negocio está aislada de la infraestructura.',
    'Frontend más limpio: ui_tweak.py pasó de 258 líneas (~15 responsabilidades) a una capa delgada.',
    'Sin regresión: Todos los endpoints y funcionalidades se mantienen exactamente igual.',
]
for b in benefits:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('— Fin de la documentación —')
run.italic = True
run.font.color.rgb = RGBColor(150, 150, 150)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Documentacion_Refactorizacion_FinTrack.docx')
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
